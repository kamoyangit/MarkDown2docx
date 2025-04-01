import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import re

def apply_formatting(paragraph, text):
    """マークダウンの装飾をWordの書式に変換"""
    # 太字 (**text**) の処理
    bold_parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(bold_parts):
        if i % 2 == 1:  # **で囲まれた部分
            run = paragraph.add_run(part)
            run.bold = True
        else:
            paragraph.add_run(part)
    
    # 斜体 (*text*) の処理
    formatted_text = paragraph.text
    paragraph.clear()
    italic_parts = re.split(r'\*(.*?)\*', ''.join(bold_parts))
    for i, part in enumerate(italic_parts):
        if i % 2 == 1:  # *で囲まれた部分
            run = paragraph.add_run(part)
            run.italic = True
        else:
            paragraph.add_run(part)

def markdown_to_word(markdown_text):
    doc = Document()
    
    # スタイル設定（オプション）
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Meiryo'  # メイリオを使用（日本語対応）
    font.size = Pt(10)
    
    for line in markdown_text.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        # 見出し処理
        if line.startswith('#'):
            level = line.count('#', 0, 6)  # 最大6レベルまで
            heading_text = line.lstrip('#').strip()
            doc.add_heading(heading_text, level=level-1)
        
        # リスト処理
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            apply_formatting(p, line[2:].strip())
        
        # 番号付きリスト
        elif re.match(r'^\d+\. ', line):
            p = doc.add_paragraph(style='List Number')
            apply_formatting(p, line[line.find(' ')+1:].strip())
        
        # 通常の段落
        else:
            p = doc.add_paragraph()
            apply_formatting(p, line)
    
    return doc

def save_docx(doc):
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# Streamlit UI
st.title('マークダウンからWordへ変換（装飾対応版）')
st.markdown("""
以下のマークダウン記法が使用できます:
- `# 見出し` → 見出し
- `**太字**` → 太字
- `*斜体*` → 斜体
- `- リスト` → 箇条書き
- `1. アイテム` → 番号付きリスト
""")

sample_md = """# サンプルドキュメント

これは*斜体*と**太字**を**組み*合*わせ**た例です。

## 特徴リスト

- **主要機能**
  - *高速*処理
  - マルチ**プラットフォーム**対応
- 追加機能
  - カスタマイズ可能

1. **第一の**ステップ
2. *第二の*ステップ
3. 最終ステップ
"""

markdown_text = st.text_area(
    'マークダウンテキストを入力してください',
    height=300,
    value=sample_md
)

if st.button('Wordに変換'):
    if markdown_text:
        try:
            doc = markdown_to_word(markdown_text)
            word_file = save_docx(doc)
            
            st.success('変換完了！')
            st.download_button(
                label='Wordファイルをダウンロード',
                data=word_file,
                file_name='formatted_document.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        except Exception as e:
            st.error(f'エラーが発生しました: {str(e)}')
    else:
        st.warning('マークダウンテキストを入力してください')